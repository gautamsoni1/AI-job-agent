from datetime import datetime
from pathlib import Path

import structlog
from jinja2 import Environment, FileSystemLoader

logger = structlog.get_logger()

TEMPLATES_DIR = Path("app/templates")
OUTPUT_DIR = Path("generated_resumes")


class ResumeGeneratorService:
    def __init__(self):
        OUTPUT_DIR.mkdir(exist_ok=True)
        self.jinja_env = Environment(
            loader=FileSystemLoader(str(TEMPLATES_DIR)),
            autoescape=True,
        )

    async def generate_pdf(self, resume_data: dict, template: str = "resume_ats_clean", user_id: str = "") -> str:
        """Generate PDF resume from structured data. Returns file path."""
        html_content = self._render_template(f"{template}.html", resume_data)
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        filename = f"resume_{user_id}_{timestamp}.pdf"
        output_path = OUTPUT_DIR / filename

        try:
            from weasyprint import HTML
            HTML(string=html_content, base_url=str(TEMPLATES_DIR)).write_pdf(str(output_path))
            logger.info("resume_pdf_generated", path=str(output_path))
            return str(output_path)
        except Exception as e:
            logger.warning("weasyprint_pdf_failed_using_fallback", error=str(e))
            return self._generate_pdf_with_pymupdf(resume_data, output_path)

    def _generate_pdf_with_pymupdf(self, resume_data: dict, output_path: Path) -> str:
        """Pure-Python/native-light PDF fallback using PyMuPDF, already used by the parser."""
        import fitz
        import textwrap

        doc = fitz.open()
        page = doc.new_page(width=595, height=842)
        x = 54
        y = 48
        line_height = 12
        page_bottom = 790

        def ensure_space(lines: int = 1):
            nonlocal page, y
            if y + (lines * line_height) > page_bottom:
                page = doc.new_page(width=595, height=842)
                y = 48

        def write(text: str, size: int = 10, bold: bool = False, indent: int = 0):
            nonlocal y
            if not text:
                return
            font = "helv-bold" if bold else "helv"
            width = 82 if indent else 92
            for line in textwrap.wrap(str(text), width=width) or [""]:
                ensure_space()
                page.insert_text((x + indent, y), line, fontsize=size, fontname=font, color=(0, 0, 0))
                y += line_height

        def section(title: str):
            nonlocal y
            y += 6
            write(title.upper(), size=10, bold=True)
            y += 2

        write(resume_data.get("full_name", "Candidate"), size=18, bold=True)
        contact = resume_data.get("contact", {}) or {}
        contact_line = " | ".join(
            str(contact.get(key, "")).strip()
            for key in ("email", "phone", "linkedin", "location")
            if contact.get(key)
        )
        write(contact_line, size=9)

        if resume_data.get("summary"):
            section("Professional Summary")
            write(resume_data["summary"])

        if resume_data.get("skills"):
            section("Skills")
            write(", ".join(str(skill) for skill in resume_data["skills"]))

        if resume_data.get("experience"):
            section("Work Experience")
            for exp in resume_data["experience"]:
                write(" - ".join(filter(None, [exp.get("title", ""), exp.get("company", "")])), bold=True)
                meta = " | ".join(filter(None, [
                    " - ".join(filter(None, [exp.get("start_date", ""), exp.get("end_date", "")])),
                    exp.get("location", ""),
                ]))
                write(meta, size=9)
                for bullet in exp.get("bullets", []):
                    write(f"- {bullet}", indent=12)

        if resume_data.get("projects"):
            section("Projects")
            for project in resume_data["projects"]:
                write(project.get("name", "Project"), bold=True)
                write(project.get("description", ""))
                if project.get("technologies"):
                    write("Technologies: " + ", ".join(project["technologies"]))

        if resume_data.get("education"):
            section("Education")
            for edu in resume_data["education"]:
                write(" - ".join(filter(None, [edu.get("degree", ""), edu.get("institution", ""), edu.get("year", "")])), bold=True)

        if resume_data.get("certifications"):
            section("Certifications")
            for certification in resume_data["certifications"]:
                write(str(certification))

        doc.save(str(output_path))
        doc.close()
        logger.info("resume_pdf_generated_fallback", path=str(output_path))
        return str(output_path)

    async def generate_docx(self, resume_data: dict, user_id: str = "") -> str:
        """Generate an ATS-friendly DOCX resume. Returns file path."""
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt, RGBColor

        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.55)
        section.bottom_margin = Inches(0.55)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

        styles = doc.styles
        styles["Normal"].font.name = "Arial"
        styles["Normal"].font.size = Pt(10)

        def add_section(title: str):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(7)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(title.upper())
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.color.rgb = RGBColor(31, 78, 121)
            p.add_run("\n")
            border = p._p.get_or_add_pPr()
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            pbdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "1F4E79")
            pbdr.append(bottom)
            border.append(pbdr)

        name = resume_data.get("full_name", "")
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(1)
        run = title_para.add_run(name)
        run.bold = True
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(31, 78, 121)

        contact = resume_data.get("contact", {})
        contact_line = " | ".join(filter(None, [
            contact.get("email"),
            contact.get("phone"),
            contact.get("linkedin"),
            contact.get("location"),
        ]))
        if contact_line:
            add_section("Contact")
            contact_para = doc.add_paragraph(contact_line)
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.paragraph_format.space_after = Pt(8)

        if resume_data.get("summary"):
            add_section("Professional Summary")
            doc.add_paragraph(resume_data["summary"])

        if resume_data.get("skills"):
            add_section("Skills")
            doc.add_paragraph(", ".join(str(skill) for skill in resume_data["skills"]))

        if resume_data.get("experience"):
            add_section("Work Experience")
            for exp in resume_data["experience"]:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(1)
                heading = " - ".join(filter(None, [exp.get("title", ""), exp.get("company", "")]))
                run = p.add_run(heading)
                run.bold = True

                meta = " | ".join(filter(None, [
                    " - ".join(filter(None, [exp.get("start_date", ""), exp.get("end_date", "Present")])),
                    exp.get("location", ""),
                ]))
                if meta.strip():
                    meta_p = doc.add_paragraph(meta)
                    meta_p.paragraph_format.space_after = Pt(2)
                for bullet in exp.get("bullets", []):
                    bullet_p = doc.add_paragraph(str(bullet), style="List Bullet")
                    bullet_p.paragraph_format.space_after = Pt(1)

        if resume_data.get("projects"):
            add_section("Projects")
            for project in resume_data["projects"]:
                p = doc.add_paragraph()
                run = p.add_run(project.get("name", "Project"))
                run.bold = True
                if project.get("description"):
                    doc.add_paragraph(project["description"])
                if project.get("technologies"):
                    doc.add_paragraph("Technologies: " + ", ".join(project["technologies"]))

        if resume_data.get("education"):
            add_section("Education")
            for edu in resume_data["education"]:
                p = doc.add_paragraph()
                run = p.add_run(" - ".join(filter(None, [edu.get("degree", ""), edu.get("institution", "")])))
                run.bold = True
                if edu.get("year"):
                    doc.add_paragraph(str(edu["year"]))

        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        filename = f"resume_{user_id}_{timestamp}.docx"
        output_path = OUTPUT_DIR / filename
        doc.save(str(output_path))
        logger.info("resume_docx_generated", path=str(output_path))
        return str(output_path)

    def _render_template(self, template_name: str, data: dict) -> str:
        try:
            template = self.jinja_env.get_template(template_name)
            return template.render(**data)
        except Exception as e:
            logger.error("template_render_failed", template=template_name, error=str(e))
            raise
