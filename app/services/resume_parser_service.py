import io
import re
from pathlib import Path
from typing import Optional

import pdfplumber
import fitz  # PyMuPDF
import structlog
from docx import Document

logger = structlog.get_logger()


class ResumeParserService:
    """Extracts raw text and structured data from PDF/DOCX resumes."""

    async def parse(self, file_path: str, file_type: str) -> dict:
        """Parse resume file and return structured data."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if file_type in ("pdf", "application/pdf"):
            raw_text = await self._extract_pdf_text(file_path)
        elif file_type in ("docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
            raw_text = await self._extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

        if not raw_text or len(raw_text.strip()) < 50:
            raise ValueError("Could not extract meaningful text from resume")

        structured = self._extract_structure(raw_text)
        return {
            "raw_text": raw_text,
            "word_count": len(raw_text.split()),
            "char_count": len(raw_text),
            **structured,
        }

    async def _extract_pdf_text(self, file_path: str) -> str:
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            logger.warning("pdfplumber_failed", error=str(e), file=file_path)

        if len(text.strip()) < 50:
            try:
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text() + "\n"
                doc.close()
            except Exception as e:
                logger.error("pymupdf_failed", error=str(e), file=file_path)

        return text.strip()

    async def _extract_docx_text(self, file_path: str) -> str:
        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            return "\n".join(paragraphs)
        except Exception as e:
            logger.error("docx_parse_failed", error=str(e), file=file_path)
            raise

    def _extract_structure(self, text: str) -> dict:
        return {
            "contact_info": self._extract_contact(text),
            "sections": self._identify_sections(text),
            "skills_found": self._extract_skills_list(text),
            "education_found": self._has_education(text),
            "experience_found": self._has_experience(text),
        }

    def _extract_contact(self, text: str) -> dict:
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        phone_pattern = r'(\+?\d{1,3}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,9})'
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        github_pattern = r'github\.com/[\w-]+'

        emails = re.findall(email_pattern, text)
        phones = re.findall(phone_pattern, text)
        linkedin = re.findall(linkedin_pattern, text, re.IGNORECASE)
        github = re.findall(github_pattern, text, re.IGNORECASE)

        return {
            "email": emails[0] if emails else None,
            "phone": phones[0] if phones else None,
            "linkedin": linkedin[0] if linkedin else None,
            "github": github[0] if github else None,
        }

    def _identify_sections(self, text: str) -> list[str]:
        section_keywords = [
            "experience", "work experience", "employment", "education", "skills",
            "projects", "certifications", "summary", "objective", "achievements",
            "awards", "publications", "languages", "interests", "volunteering"
        ]
        found = []
        text_lower = text.lower()
        for kw in section_keywords:
            if kw in text_lower:
                found.append(kw.title())
        return found

    def _extract_skills_list(self, text: str) -> list[str]:
        common_skills = [
            "Python", "JavaScript", "TypeScript", "Java", "C++", "C#", "Go", "Rust",
            "FastAPI", "Django", "Flask", "React", "Vue", "Angular", "Node.js",
            "MongoDB", "PostgreSQL", "MySQL", "Redis", "Elasticsearch",
            "Docker", "Kubernetes", "AWS", "GCP", "Azure", "Terraform",
            "Machine Learning", "Deep Learning", "TensorFlow", "PyTorch",
            "Git", "CI/CD", "REST API", "GraphQL", "Microservices",
            "SQL", "NoSQL", "Linux", "Agile", "Scrum"
        ]
        found = []
        for skill in common_skills:
            if skill.lower() in text.lower():
                found.append(skill)
        return found

    def _has_education(self, text: str) -> bool:
        education_keywords = ["bachelor", "master", "phd", "degree", "university", "college", "b.tech", "m.tech", "b.e", "m.e"]
        text_lower = text.lower()
        return any(kw in text_lower for kw in education_keywords)

    def _has_experience(self, text: str) -> bool:
        experience_keywords = ["experience", "worked at", "employed", "position", "role", "job"]
        text_lower = text.lower()
        return any(kw in text_lower for kw in experience_keywords)